"""
Basic Encryption & Decryption Program

This Python program was created in 2021 as a personal project to
experiment with basic encryption and decryption techniques.

The program uses a custom character-to-binary mapping to encode
text from a Word document into a binary representation and provides
the ability to reverse the process and recover the original message.

The user can select whether to encrypt or decrypt a document and
specify the input and output file locations.

Built with:
- Python
- python-docx
"""

import docx
import os
import time

dict = {'v':'00100000' , '!':'00100001' , '"':'00100010' , '2':'00100011' , ',':'00100100' , '4':'00100101' , '.':'00100110' ,
        '/':'00101000' , ')':'00101001' , 'c':'00101010' , '9':'00101011' , '$':'00101100' , ';':'00101101' , '&':'00101110' ,
        '(':'00101111' , '>':'00110000' , '?':'00110001' , '#':'00110010' , 'H':'00110011' , '%':'00110100' , 'Q':'00110101' ,
        '6':'00110110' , '7':'00110111' , '8':'00111000' , '+':'00111001' , ':':'00111010' , '-':'00111011' , 'C':'00111100' ,
        '=':'00111101' , '0':'00111110' , '1':'00111111' , 'U':'01000000' , 'e':'01000001' , 't':'01000010' , '<':'01000011' ,
        'K':'01000100' , 'S':'01000101' , '[':'01000110' , 'r':'01000111' , '3':'01001000' , 'P':'01001001' , 'J':'01001010' ,
        'D':'01001011' , 'Z':'01001100' , 'M':'01001101' , 'd':'01001110' , 'V':'01001111' , 'I':'01010000' , '5':'01010001' ,
        'R':'01010010' , 'E':'01010011' , '*':'01010100' , '@':'01010101' , 'O':'01010110' , '^':'01010111' , 'n':'01011000' ,
        'a':'01011001' , 'L':'01011010' , 'F':'01011011' , 's':'01011101' , 'W':'01011110' , 'g':'01011111' , 'Y':'01100001' ,
        'p':'01100010' , 'q':'01100011' , 'N':'01100100' , 'A':'01100101' , 'B':'01100110' , '_':'01100111' , 'o':'01101000' ,
        'w':'01101001' , 'x':'01101010' , 'y':'01101011' , 'z':'01101100' , '{':'01101101' , 'X':'01101110' , 'h':'01101111' ,
        'b':'01110000' , 'T':'01110001' , 'G':'01110010' , ']':'01110011' , 'f':'01110100' , '}':'01110101' , ' ':'01110110' ,
        'i':'01110111' , 'j':'01111000' , 'k':'01111001' , 'l':'01111010' , 'm':'01111011' , 'u':'01111101' }



sel_num = int(input("To encrypt - Press 1\nTo Decrypt - Press 2\nYour Selection: "))
match sel_num:
    case 1:
        # Input the address of the file to be encrypted
        in_address = input("Please Enter the address of the file to be encrypted: ")
        in_address = in_address.replace("\\","/") #this changes the bck slash to forward slash

        # input the name of the file with the extension
        file_name = input("Please Enter the name of the file(with extension)to be encrypted: ")
        file_name = file_name.replace("\\","/") #this changes the bck slash to forward slash

        # complete address
        comp_address = in_address + "/" + file_name

        print("Encrytion initiating")
        time.sleep(2)
        print("Reading the Message..")
        time.sleep(2)

        #reading the original message
        rdoc = docx.Document(comp_address)
        completeText = []
        for paragraph in rdoc.paragraphs:
                completeText.append(paragraph.text)

        completeText = [item.lower() for item in completeText]  # all items in list lowercase

        full_str = completeText[0]  #from paragraph to str
        num = 0
        encd_str = ''


        print("Encrypting...")
        time.sleep(5)

        while(num<len(full_str)):
             orig_str = full_str[num]
             num = num + 1
             temp_str = dict.get(orig_str)
             encd_str = str(encd_str) + str(temp_str)


        print("Message Encrypted.!")
        time.sleep(2)

        # Input the address to save
        out_address = input("Please Enter the address where to save: ")
        out_address = out_address.replace("\\","/") #this changes the bck slash to forward slash

        # the name of the file
        out_file_name = out_address + "/Encrypted_Message.docx"  
        #writing the encoded message now
        encded_Doc = docx.Document()
        parag = encded_Doc.add_paragraph(encd_str)
        encded_Doc.save(out_file_name)

        

        op_num = int(input("To open the file press 1\nPress 2 to exit\nYour selection: "))
        if op_num == 1:
            print("Opening the Encrypted file..")
            time.sleep(2)
            os.system(out_file_name)
        else:
              exit
    
    case 2:
          
                def get_key(val):
                     for key, value in dict.items():
                          if val == value:
                             return key


                # Input the address of the file to be decrypted
                in_address = input("Please Enter the address of the file to be decrypted: ")
                in_address = in_address.replace("\\","/") #this changes the bck slash to forward slash

                # input the name of the file with the extension
                file_name = input("Please Enter the name of the file(with extension) to be encrypted: ")
                file_name = file_name.replace("\\","/") #this changes the bck slash to forward slash

                # complete address
                comp_address = in_address + "/" + file_name


                print("Reading the Message..")
                time.sleep(2)


                #reading the original message
                rdoc = docx.Document(comp_address)
                completeText = []
                for paragraph in rdoc.paragraphs:
                        completeText.append(paragraph.text)


                full_str = completeText[0]  #from paragraph to str
                num = 0
                decrypted_str = ''
                leng = len(full_str)

                print("Decrypting...")
                time.sleep(5)

                while(num<leng):
                      temp_str = full_str[num:num + 8]
                      num += 8
                      temp_str2 = str(get_key(temp_str))
                      decrypted_str += temp_str2

                print("Message Decrypted.!")
                time.sleep(2)


                # Input the address to save
                out_address = input("Please Enter the address where to save: ")
                out_address = out_address.replace("\\","/") #this changes the bck slash to forward slash

                # the name of the file
                out_file_name = out_address + "/Decrypted_Message.docx"  

                #writing the decoded message now
                encded_Doc = docx.Document()
                parag = encded_Doc.add_paragraph(decrypted_str)
                encded_Doc.save(out_file_name)


                

                op_num = int(input("To open the file press 1\nPress 2 to exit\nYour selection: "))
                if op_num == 1:
                    print("Opening the Decrypted file..")
                    time.sleep(2)
                    os.system(out_file_name)
                else:
                    exit
    case _ :  print("wrong selection")
print("Thankyou for using me !")